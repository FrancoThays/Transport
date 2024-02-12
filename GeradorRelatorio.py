'''
objeto para produção em relatório de docx
Thays Franco
2022-04-08
'''
#================================================
#IMPORTS

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches

#================================================

class GeradorRelatorio:

    def __init__(self):
        self.document = Document()
        self.document.add_heading("Relatório de processamento de produção de viagens",0)
        self.document.add_paragraph("Este é o primeiro relatório referente ao projeto 1234.567/JLM. Ele trata do processamento de dados de demanda referentes a produção e a atração de viagens e aos dados socioeconomicos...")

    def insere_secao(self, titulo, secao):
        self.document.add_heading(titulo, secao)

    def insere_paragrafo(self, paragrafo):
        self.document.add_paragraph(paragrafo)

    def insere_tabela(self, df):
        table = self.document.add_table(df.shape[0]+1, df.shape[1])
        table.style = 'Table Grid'

        #adicionando nome das colunas
        for j in range(df.shape[1]):
            table.cell(0, j).text = df.columns[j]

        #adicionando os dados
        for i in range(df.shape[0]):
            for j in range(df.shape[1]):
                table.cell(i+1, j).text = str(round(df.values[i, j]))
                table.cell(i+1,j).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.RIGHT


    def insere_quebra_de_pagina(self):
        self.document.add_page_break()

    def insere_imagem(self, grafico):
        self.document.add_picture("../RESULTADOS/grafico.png")

    def conclui_documento(self, out_file_name):
        self.document.save(out_file_name)

